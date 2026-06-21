from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\jvlce\Documents\namoro\Roteiro_Dungeon_Escape_Simples.docx"


def set_run_font(run, name="Calibri"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def set_style_font(style, name="Calibri", size=11, color=None):
    style.font.name = name
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11, "000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, "Calibri", 16, "2E74B5")
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.25

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, "Calibri", 13, "2E74B5")
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.25

    bullet = doc.styles["List Bullet"]
    set_style_font(bullet, "Calibri", 11, "000000")
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Roteiro simples - Dungeon Escape")
    set_run_font(run)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("000000")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Comentários curtos para explicar para que serve cada parte do código")
    set_run_font(run)
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("555555")


def add_body_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run)
    run.font.size = Pt(11)
    return p


def add_comment_block(doc, title, bullets):
    doc.add_heading(title, level=2)
    for bullet in bullets:
        add_bullet(doc, bullet)


doc = Document()
configure_styles(doc)
add_title(doc)

add_body_paragraph(
    doc,
    "Use este roteiro como uma cola. Abra o projeto no VS Code, mostre os arquivos principais "
    "e depois desça o sketch.js de cima para baixo. A ideia não é explicar linha por linha, "
    "e sim comentar o objetivo de cada bloco do código.",
)

doc.add_heading("Antes de entrar no sketch.js", level=1)
add_comment_block(
    doc,
    "Arquivos do projeto",
    [
        "index.html: carrega a página do jogo, a biblioteca p5.js, o p5.sound e o arquivo sketch.js.",
        "style.css: deixa a página organizada e centraliza a área do jogo na tela.",
        "assets/: guarda as imagens e sons usados no jogo, como jogador, inimigo, chave, porta, armadilha, efeitos e música.",
        "README.md: explica o objetivo do jogo, os controles e como executar o projeto.",
    ],
)

doc.add_heading("Comentários simples para o sketch.js", level=1)

sections = [
    (
        "1. Constantes, assets e fases - início do arquivo",
        [
            "CANVAS_WIDTH, CANVAS_HEIGHT e HUD_HEIGHT: definem o tamanho da tela e o espaço do painel de informações.",
            "STARTING_LIVES e INVINCIBLE_TIME: definem as vidas iniciais e o tempo em que o jogador fica protegido depois de tomar dano.",
            "assetFiles: junta os caminhos das imagens e sons em um só lugar, para facilitar a manutenção.",
            "levels: guarda as fases do jogo. Cada fase tem nome, tempo, velocidade dos inimigos e um mapa escrito com caracteres.",
            "No mapa, # é parede, . é chão, P é jogador, K é chave, D é porta, E é inimigo e T é armadilha.",
        ],
    ),
    (
        "2. preload(), setup() e draw()",
        [
            "preload(): carrega imagens e sons antes do jogo começar, para evitar erro durante a partida.",
            "setup(): cria o canvas, coloca ele na página e cria o GameManager, que controla o jogo.",
            "draw(): é o loop principal do p5.js. Ele roda várias vezes por segundo e pede para o GameManager atualizar e desenhar tudo.",
        ],
    ),
    (
        "3. Mouse, teclado e som",
        [
            "mousePressed(): libera o áudio do navegador e verifica se algum botão do menu foi clicado.",
            "keyPressed(): trata teclas importantes, como espaço para pausar, Enter para iniciar, R para reiniciar e ESC para voltar.",
            "unlockAudio(): libera o som, porque o navegador normalmente só permite áudio depois de uma interação do usuário.",
            "playSound(): toca efeitos sonoros, como pegar chave, tomar dano, vencer ou perder.",
        ],
    ),
    (
        "4. Música e pause",
        [
            "startGameMusic(): inicia a música de fundo durante a fase.",
            "pauseGameMusic(): pausa a música quando o jogo é pausado.",
            "stopGameMusic(): para a música quando o jogador volta ao menu, perde ou vence.",
            "togglePause(): alterna entre jogo pausado e jogo rodando.",
            "drawPauseOverlay(): desenha a camada escura com a mensagem de jogo pausado.",
        ],
    ),
    (
        "5. Funções auxiliares",
        [
            "collidesAabb(): verifica colisão retangular entre dois objetos, como jogador e chave ou jogador e inimigo.",
            "drawPanel(): desenha painéis usados nos menus e telas finais.",
            "drawCenteredTitle(): centraliza títulos como Dungeon Escape, Game Over e Vitória.",
        ],
    ),
    (
        "6. Classe GameManager",
        [
            "GameManager é a classe principal do jogo. Ela guarda o estado atual, fase, vidas, score, tempo, botões e objetos.",
            "setupMenu(), setupInstructions(), setupAbout(), setupGameOver() e setupVictory(): montam as telas do jogo.",
            "startGame(): zera a pontuação, reinicia as vidas e começa a primeira fase.",
            "loadLevel(): transforma o mapa de texto em objetos reais do jogo, criando jogador, chave, porta, inimigos e armadilhas.",
        ],
    ),
    (
        "7. Atualização da partida",
        [
            "update(): decide o que deve acontecer dependendo do estado do jogo, como menu, fase, pausa, game over ou vitória.",
            "updatePlaying(): atualiza o jogador, os inimigos, as partículas e as colisões durante a fase.",
            "checkItemCollision(): verifica se o jogador pegou a chave e soma pontos.",
            "checkDoorCollision(): verifica se o jogador chegou na porta com a chave para passar de fase.",
            "checkDamageCollision(): verifica se o jogador encostou em inimigo ou armadilha.",
            "damagePlayer(): tira vida, tira pontos e deixa o jogador invencível por pouco tempo.",
            "completeLevel(): soma pontos da fase, calcula bônus de tempo e passa para a próxima fase ou mostra a vitória.",
        ],
    ),
    (
        "8. Tempo, mapa e paredes",
        [
            "calculateTimeLeft(): calcula quanto tempo ainda resta na fase.",
            "calculateTimeBonus(): transforma o tempo restante em pontos extras.",
            "tileCenter(): converte linha e coluna do mapa em uma posição real dentro do canvas.",
            "isWallAt(): verifica se uma posição é parede ou está fora do mapa.",
            "collidesWithWalls(): impede o jogador de atravessar paredes.",
            "collidesWithSolid(): também bloqueia a porta quando ela ainda está fechada.",
        ],
    ),
    (
        "9. Desenho das telas e do HUD",
        [
            "draw(): escolhe qual tela desenhar: menu, instruções, sobre, fase, game over ou vitória.",
            "drawGameScene(): desenha a cena da fase, incluindo mapa, objetos, jogador, HUD e mensagens.",
            "drawWorld(): percorre o mapa e desenha chão ou parede em cada quadrado.",
            "drawHud(): mostra fase, vidas, score, tempo e se a chave já foi coletada.",
            "drawGameOver() e drawVictory(): desenham as telas finais com a pontuação.",
        ],
    ),
    (
        "10. Classe Player",
        [
            "Player representa o personagem controlado pelo jogador.",
            "update(): lê WASD ou setas e calcula a direção do movimento.",
            "tryMove(): tenta mover o jogador, mas só muda a posição se não houver colisão.",
            "makeInvincible() e isInvincible(): controlam a proteção temporária depois de tomar dano.",
            "draw(): desenha o personagem e mostra um efeito visual quando ele está com a chave.",
        ],
    ),
    (
        "11. Outras classes do jogo",
        [
            "Enemy: representa os inimigos que se movem e mudam de direção quando batem em uma parede.",
            "Item: representa a chave que o jogador precisa pegar.",
            "Door: representa a porta de saída, que só abre depois da chave.",
            "Trap: representa as armadilhas paradas que causam dano.",
            "Button: cria os botões usados nos menus.",
            "Particle: cria efeitos visuais pequenos, como brilho, dano e conclusão de fase.",
        ],
    ),
]

for title, bullets in sections:
    add_comment_block(doc, title, bullets)

doc.add_heading("Fechamento curto", level=1)
add_body_paragraph(
    doc,
    "Para encerrar, eu posso dizer que o jogo usa JavaScript com p5.js para criar o canvas, "
    "desenhar os elementos, receber teclado e mouse, tocar sons e atualizar a partida. "
    "As classes ajudam a separar responsabilidades, e o GameManager junta tudo para controlar "
    "menus, fases, vidas, pontuação, pause, game over e vitória.",
)

doc.add_heading("Frase final", level=1)
add_body_paragraph(
    doc,
    "Esse foi o Dungeon Escape. O código foi organizado em partes para facilitar a manutenção: "
    "uma parte carrega os arquivos, outra define as fases, outra controla o jogo e as classes "
    "representam os objetos principais da partida.",
)

doc.save(OUTPUT)
print(OUTPUT)
